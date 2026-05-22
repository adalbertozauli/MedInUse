import re
from dataclasses import dataclass
from pathlib import Path


@dataclass(frozen=True)
class Medication:
    name: str
    schedule: str

    def formatted(self) -> str:
        return f"{self.name} ({self.schedule})"


def read_docx_lines(path: str | Path) -> list[str]:
    from docx import Document

    document = Document(str(path))
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split())
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = " ".join(cell.text.split())
                if text:
                    lines.append(text)

    return lines


def extract_medications_from_docx(path: str | Path) -> list[Medication]:
    return extract_medications(read_docx_lines(path))


def extract_medications(lines: list[str]) -> list[Medication]:
    medications: list[Medication] = []
    index = 0

    while index < len(lines):
        current = lines[index].strip()
        next_line = lines[index + 1].strip() if index + 1 < len(lines) else ""

        if _looks_like_medication(current) and _looks_like_posology(next_line):
            medications.append(Medication(current, infer_schedule(next_line)))
            index += 2
            continue

        if _looks_like_posology(current) and medications:
            index += 1
            continue

        index += 1

    return medications


def format_medications(medications: list[Medication], one_per_line: bool) -> str:
    items = [medication.formatted() for medication in medications]
    if one_per_line:
        return "\n".join(f"- {item}" for item in items)
    return "; ".join(items)


def infer_schedule(posology: str) -> str:
    text = _normalize(posology)
    quantity = _dose_quantity(text)

    long_interval_schedule = _long_interval_schedule(text, quantity)
    if long_interval_schedule:
        return long_interval_schedule

    weekly_schedule = _weekly_schedule(text)
    if weekly_schedule:
        return weekly_schedule

    unique_dose = _unique_dose_schedule(text, quantity)
    if unique_dose:
        return unique_dose

    interval_match = re.search(r"(\d{1,2})\s*/\s*(\d{1,2})\s*h(?:oras?)?", text)
    if interval_match:
        first, second = interval_match.groups()
        if first == second:
            return _with_quantity(_schedule_for_hour_interval(int(first)), quantity)

    hour_match = re.search(r"de\s+(\d{1,2})\s+em\s+(\d{1,2})\s+h(?:oras?)?", text)
    if hour_match and hour_match.group(1) == hour_match.group(2):
        return _with_quantity(_schedule_for_hour_interval(int(hour_match.group(1))), quantity)

    every_hour_match = re.search(r"(?:a\s+cada|cada|apos|depois\s+de)\s+(\d{1,2})\s*h(?:oras?)?", text)
    if every_hour_match:
        return _with_quantity(_schedule_for_hour_interval(int(every_hour_match.group(1))), quantity)

    times_per_day = _times_per_day(text)
    if times_per_day:
        return _with_quantity(_schedule_for_daily_frequency(times_per_day), quantity)

    if any(term in text for term in ("6/6", "seis em seis")):
        return _with_quantity("1-1-1-1", quantity)
    if any(term in text for term in ("8/8", "oito em oito")):
        return _with_quantity("1-1-1", quantity)
    if any(term in text for term in ("12/12", "doze em doze")):
        return _with_quantity("1-0-1", quantity)

    if _mentions_morning(text) and _mentions_afternoon(text) and _mentions_night(text):
        return _with_quantity("1-1-1", quantity)
    if _mentions_all_meals(text):
        return _with_quantity("1-1-1", quantity)
    if _mentions_morning(text) and _mentions_night(text):
        return _with_quantity("1-0-1", quantity)
    if _mentions_morning(text) and _mentions_afternoon(text):
        return _with_quantity("1-1-0", quantity)
    if _mentions_afternoon(text) and _mentions_night(text):
        return _with_quantity("0-1-1", quantity)

    time_schedule = _schedule_for_explicit_times(text)
    if time_schedule:
        return _with_quantity(time_schedule, quantity)

    if _mentions_morning(text):
        return _with_quantity("1-0-0", quantity)
    if _mentions_afternoon(text):
        return _with_quantity("0-1-0", quantity)
    if _mentions_night(text):
        return _with_quantity("0-0-1", quantity)

    if _mentions_once_daily(text):
        return _with_quantity("1x/dia", quantity)

    return "conferir posologia"


def _looks_like_medication(text: str) -> bool:
    normalized = _normalize(text)
    if _is_section_or_admin_line(normalized):
        return False
    if not text or _looks_like_posology(text):
        return False
    has_dosage = re.search(r"\d+(?:[,.]\d+)?\s*(?:mg|mcg|g|ml|ui|%)\b", normalized) is not None
    return has_dosage or len(text.split()) <= 5


def _looks_like_posology(text: str) -> bool:
    normalized = _normalize(text)
    return any(
        term in normalized
        for term in (
            "tomar",
            "usar",
            "aplicar",
            "ingerir",
            "aspirar",
            "inalar",
            "dar",
            "12/12",
            "8/8",
            "6/6",
            "a cada",
            "ao dia",
            "por dia",
            "vez ao dia",
            "vezes ao dia",
            "diario",
            "dose unica",
            "refeicoes",
            "refeicao",
            "manha",
            "cedo",
            "almoco",
            "jantar",
            "noite",
            "jejum",
            "deitar",
        )
    )


def _schedule_for_hour_interval(hours: int) -> str:
    schedules = {
        6: "1-1-1-1",
        8: "1-1-1",
        12: "1-0-1",
        24: "1-0-0",
    }
    return schedules.get(hours, "conferir posologia")


def _schedule_for_daily_frequency(times: int) -> str:
    schedules = {
        1: "1x/dia",
        2: "1-0-1",
        3: "1-1-1",
        4: "1-1-1-1",
    }
    return schedules.get(times, "conferir posologia")


def _weekly_schedule(text: str) -> str | None:
    if not _mentions_weekly(text):
        return None

    quantity = _dose_quantity(text) or "1 comp."
    duration = _duration_in_weeks(text)

    if duration:
        return f"{quantity} semana/{duration} semanas"
    return f"{quantity} semana"


def _mentions_weekly(text: str) -> bool:
    return any(
        term in text
        for term in (
            "por semana",
            "uma vez por semana",
            "1 vez por semana",
            "1x por semana",
            "1x na semana",
            "1x/semana",
            "semanal",
            "semanalmente",
            "a cada semana",
        )
    )


def _dose_quantity(text: str) -> str:
    number_pattern = r"(\d{1,3}(?:[,.]\d+)?|uma|duas|dois|tres|quatro|meio|meia|½)"
    unit_patterns = (
        (rf"{number_pattern}\s*(?:comprimido|comprimidos|comp\b|capsula|capsulas)", "comp.", "comps."),
        (rf"{number_pattern}\s*(?:gota|gotas|gts?\b)", "gota", "gotas"),
        (rf"{number_pattern}\s*(?:ml|m l|mililitro|mililitros)", "ml", "ml"),
        (rf"{number_pattern}\s*(?:unidade|unidades|u\b|ui\b)", "U", "U"),
        (rf"{number_pattern}\s*(?:puff|puffs|jato|jatos)", "puff", "puffs"),
        (rf"{number_pattern}\s*(?:ampola|ampolas|amp\.?)", "ampola", "ampolas"),
        (rf"{number_pattern}\s*(?:dose|doses)", "dose", "doses"),
        (rf"{number_pattern}\s*(?:colher|colheres)", "colher", "colheres"),
    )

    for pattern, singular, plural in unit_patterns:
        match = re.search(pattern, text)
        if not match:
            continue

        raw_amount = match.group(1)
        amount = _number_to_float(raw_amount)
        if amount is None:
            continue

        unit = singular if amount <= 1 else plural
        return f"{_format_amount(amount)} {unit}"

    return ""


def _with_quantity(schedule: str, quantity: str) -> str:
    if not quantity or quantity in ("1 comp.", "1 comps."):
        return schedule
    expanded = _expanded_dose_schedule(schedule, quantity)
    if expanded:
        return expanded
    return f"{quantity} {schedule}"


def _long_interval_schedule(text: str, quantity: str) -> str:
    interval_match = re.search(r"(\d{1,3})\s*/\s*(\d{1,3})\s*dias?", text)
    if interval_match and interval_match.group(1) == interval_match.group(2):
        return _format_long_interval(quantity, int(interval_match.group(1)))

    every_days_match = re.search(r"(?:a\s+cada|cada|de\s+)\s*(\d{1,3})\s*dias?", text)
    if every_days_match:
        return _format_long_interval(quantity, int(every_days_match.group(1)))

    return ""


def _format_long_interval(quantity: str, days: int) -> str:
    dose = quantity or "1 dose"
    return f"{dose}/{days} dias"


def _unique_dose_schedule(text: str, quantity: str) -> str:
    if not any(term in text for term in ("dose unica", "dose unico")):
        return ""
    return f"{quantity or '1 dose'} dose única"


def _expanded_dose_schedule(schedule: str, quantity: str) -> str:
    if not re.fullmatch(r"[01](?:-[01])+", schedule):
        return ""

    match = re.fullmatch(r"(.+?)\s+(gota|gotas|ml|U|puff|puffs)", quantity)
    if not match:
        return ""

    amount, unit = match.groups()
    slots = schedule.split("-")

    if unit in ("gota", "gotas"):
        first_dose = f"{amount} gts"
        dose = f"{amount}gts"
    elif unit == "ml":
        first_dose = dose = f"{amount}ml"
    else:
        suffix = "U" if unit == "U" else "puffs"
        first_dose = dose = f"{amount}{suffix}"

    rendered = []
    used_first_dose = False
    for slot in slots:
        if slot == "0":
            rendered.append("0")
        elif not used_first_dose:
            rendered.append(first_dose)
            used_first_dose = True
        else:
            rendered.append(dose)

    return "-".join(rendered)


def _duration_in_weeks(text: str) -> int | None:
    week_match = re.search(r"(?:por|durante|por ate)?\s*(\d{1,2})\s*semanas?", text)
    if week_match:
        return int(week_match.group(1))

    month_match = re.search(r"(?:por|durante|por ate)?\s*(\d{1,2})\s*mes(?:es)?", text)
    if month_match:
        return int(month_match.group(1)) * 4

    return None


def _times_per_day(text: str) -> int | None:
    number_pattern = r"(\d{1,2}|uma|duas|dois|tres|quatro)"
    patterns = (
        rf"{number_pattern}\s*x\s*/?\s*(?:ao|por)?\s*dia",
        rf"{number_pattern}\s+vez(?:es)?\s+(?:ao|por)\s+dia",
        rf"{number_pattern}\s+vez(?:es)?\s+ao\s+dia",
        rf"{number_pattern}\s+vez(?:es)?\s+diaria(?:s)?",
    )

    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return _number_to_int(match.group(1))

    return None


def _number_to_int(value: str) -> int | None:
    words = {
        "uma": 1,
        "duas": 2,
        "dois": 2,
        "tres": 3,
        "quatro": 4,
        "meio": 0,
        "meia": 0,
        "½": 0,
    }
    if value.isdigit():
        return int(value)
    return words.get(value)


def _schedule_for_explicit_times(text: str) -> str:
    hours = _explicit_hours(text)
    if not hours:
        return ""

    periods = set()
    for hour in hours:
        if hour < 13:
            periods.add("morning")
        elif hour < 18:
            periods.add("afternoon")
        else:
            periods.add("night")

    if periods == {"morning"}:
        return "1-0-0"
    if periods == {"afternoon"}:
        return "0-1-0"
    if periods == {"night"}:
        return "0-0-1"
    if periods == {"morning", "afternoon"}:
        return "1-1-0"
    if periods == {"morning", "night"}:
        return "1-0-1"
    if periods == {"afternoon", "night"}:
        return "0-1-1"
    return "1-1-1"


def _explicit_hours(text: str) -> list[int]:
    hours: list[int] = []
    patterns = (
        r"(?:as|às|aos?|pelas?)\s*(\d{1,2})(?::\d{2})?\s*(?:h|horas?)?",
        r"\b(\d{1,2}):\d{2}\b",
        r"\b(\d{1,2})\s*(?:h|horas?)\b",
    )

    for pattern in patterns:
        for match in re.finditer(pattern, text):
            hour = int(match.group(1))
            if 0 <= hour <= 23 and hour not in hours:
                hours.append(hour)

    return hours


def _number_to_float(value: str) -> float | None:
    if value in ("meio", "meia", "½"):
        return 0.5
    numeric = _number_to_int(value)
    if numeric is not None:
        return float(numeric)
    try:
        return float(value.replace(",", "."))
    except ValueError:
        return None


def _format_amount(value: float) -> str:
    if value.is_integer():
        return str(int(value))
    return str(value).replace(".", ",")


def _is_section_or_admin_line(text: str) -> bool:
    clean = text.strip(" :-")
    if not clean:
        return True
    section_prefixes = (
        "uso ",
        "para",
        "paciente",
        "endereco",
        "prescricao",
        "a assistencia social",
        "à assistencia social",
        "atenciosamente",
        "cid",
    )
    return any(clean.startswith(prefix) for prefix in section_prefixes)


def _mentions_once_daily(text: str) -> bool:
    return any(
        term in text
        for term in (
            "24/24",
            "ao dia",
            "por dia",
            "diario",
            "diaria",
            "diariamente",
            "uma vez ao dia",
            "1 vez ao dia",
            "1x ao dia",
            "1x/dia",
        )
    )


def _mentions_morning(text: str) -> bool:
    return any(term in text for term in ("cedo", "manha", "cafe", "ao acordar", "jejum"))


def _mentions_afternoon(text: str) -> bool:
    return any(term in text for term in ("almoco", "tarde"))


def _mentions_night(text: str) -> bool:
    return any(term in text for term in ("noite", "jantar", "dormir", "deitar"))


def _mentions_all_meals(text: str) -> bool:
    return any(term in text for term in ("antes das refeicoes", "antes das refeicao", "antes de cada refeicao", "refeicoes"))


def _normalize(text: str) -> str:
    replacements = str.maketrans(
        {
            "á": "a",
            "à": "a",
            "ã": "a",
            "â": "a",
            "é": "e",
            "ê": "e",
            "í": "i",
            "ó": "o",
            "ô": "o",
            "õ": "o",
            "ú": "u",
            "ç": "c",
        }
    )
    return text.lower().translate(replacements)
